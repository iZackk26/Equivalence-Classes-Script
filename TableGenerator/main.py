import os
import json
import copy
from docx import Document
from docx.enum.text import WD_BREAK

from utils.gpt import generar_descripcion_y_resultados  # Función que llama a ChatGPT

# ─── 1) Cargar JSON de casos fallidos ────────────────────────────────────────
RUTA_JSON = "Data/faildSignin.json"
if not os.path.isfile(RUTA_JSON):
    raise FileNotFoundError(f"No encontré el archivo {RUTA_JSON} en el directorio actual.")

with open(RUTA_JSON, "r", encoding="utf-8") as f:
    fallidos = json.load(f)
# ────────────────────────────────────────────────────────────────────────────

# ─── 2) Ruta de la plantilla ("Table.docx") con placeholders actualizados ───
RUTA_PLANTILLA = "Tables/SignInTable.docx"
if not os.path.isfile(RUTA_PLANTILLA):
    raise FileNotFoundError(f"No encontré la plantilla {RUTA_PLANTILLA} en este directorio.")
# ────────────────────────────────────────────────────────────────────────────

# ─── 3) Crear el documento maestro vacío ────────────────────────────────────
master_doc = Document()
# ────────────────────────────────────────────────────────────────────────────

# ─── 4) Funciones auxiliares para reemplazar placeholders ──────────────────
def reemplazar_texto_en_parrafos(doc: Document, mapping: dict):
    """
    Recorre cada párrafo y reemplaza {clave} por mapping[clave].
    """
    for párrafo in doc.paragraphs:
        for clave, valor in mapping.items():
            placeholder = f"{{{clave}}}"
            if placeholder in párrafo.text:
                for run in párrafo.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(valor))


def reemplazar_texto_en_tablas(doc: Document, mapping: dict):
    """
    Recorre cada celda de cada tabla del documento y reemplaza {clave} por mapping[clave].
    """
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for clave, valor in mapping.items():
                    placeholder = f"{{{clave}}}"
                    if placeholder in celda.text:
                        for párrafo in celda.paragraphs:
                            for run in párrafo.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(valor))
# ────────────────────────────────────────────────────────────────────────────


# ─── 5) Iterar sobre cada caso y agregarlo al master_doc (offset de 34) ────
for idx, caso in enumerate(fallidos):
    # 5.1) Extraer la parte numérica de caso["codigo"] (ej. "CP-02" → 2)
    raw_code = caso["codigo"]
    parte_numerica = "".join(ch for ch in raw_code if ch.isdigit())
    if not parte_numerica:
        raise ValueError(f"No pude extraer número de '{raw_code}'")
    valor_numero = int(parte_numerica)

    # 5.2) Calcular el nuevo número: 34 + valor_numero
    nuevo_numero = 34 + valor_numero

    # 5.3) Llamar a ChatGPT para generar "resultado_esperado" y "resultado_obtenido"
    print(f"\n=== Caso {raw_code} → nuevo número {nuevo_numero} ===")
    servicio = generar_descripcion_y_resultados(caso)
    print("Respuesta completa de ChatGPT (JSON):")
    print(json.dumps(servicio, ensure_ascii=False, indent=2))

    # Extraemos los dos valores que necesitamos
    res_esperado = servicio["resultado_esperado"]
    res_obtenido = servicio["resultado_obtenido"]

    print("\n-- Resultado esperado (ChatGPT):")
    print(res_esperado)
    print("\n-- Resultado obtenido (ChatGPT):")
    print(res_obtenido)
    print("=============================================")

    # 5.4) Construir el diccionario mapping con todas las claves que aparezcan en Table.docx
    mapping = {
        # --- "numero" para CP-{numero} y Caso {numero} del Formulario SignIn ---
        "numero":             str(nuevo_numero),

        # --- Campos estáticos/manuales del JSON ---
        "is_valid":           caso["is_valid"],
        "username":           caso["username"] or "vacío",
        "password":           caso["password"] or "vacío",
        "slider":             caso["slider"],
        "remember_me":        caso["remember_me"],

        # --- "descripcion": colocamos todo en un solo placeholder ---
        "descripcion": (
            f"Probar caso {caso['is_valid']}: "
            f"Props < username {'vacío' if caso['username']=='' else caso['username']}, "
            f"password {'vacío' if caso['password']=='' else caso['password']}, "
            f"slider {caso['slider']} y checkbox “Remember Me” {caso['remember_me']} >"
        ),

        # --- "pasos": lo rellenamos justo después ---
        "pasos": "",

        # --- LOS DOS NUEVOS PLACEHOLDERS que definimos en la plantilla ---
        "res_esperado": res_esperado,
        "res_obtenido": res_obtenido,
    }

    # 5.5) Generar los 6 pasos y asignarlos a {pasos}
    paso1 = f"1. Seleccionar tipo de usuario {caso['tipo_usuario']}"
    paso2 = f"2. Ingresar “{mapping['username']}” en el campo username"
    paso3 = f"3. Ingresar “{mapping['password']}” en el campo password"
    if caso["slider"] == "arrastrado":
        paso4 = "4. Arrastrar el slider hasta el final"
    else:
        paso4 = "4. No arrastrar el slider"
    if caso["remember_me"] == "marcado":
        paso5 = '5. Marcar el checkbox “Remember Me”'
    else:
        paso5 = '5. No marcar el checkbox “Remember Me”'
    paso6 = '6. Hacer clic en “Sign In”'

    mapping["pasos"] = "\n".join([paso1, paso2, paso3, paso4, paso5, paso6])

    # 5.6) Abrir la plantilla y reemplazar los placeholders
    temp_doc = Document(RUTA_PLANTILLA)

    # Primero en párrafos
    reemplazar_texto_en_parrafos(temp_doc, mapping)
    # Luego en todas las tablas (celdas)
    reemplazar_texto_en_tablas(temp_doc, mapping)

    # 5.7) Copiar la única tabla de temp_doc al master_doc
    if not temp_doc.tables:
        raise RuntimeError("La plantilla no contiene tablas.")
    tabla_original = temp_doc.tables[0]
    tbl_xml = tabla_original._tbl
    tbl_nueva = copy.deepcopy(tbl_xml)
    master_doc._body._element.append(tbl_nueva)

    # 5.8) Insertar salto de página si no es el último caso
    if idx < len(fallidos) - 1:
        par = master_doc.add_paragraph()
        run = par.add_run()
        run.add_break(WD_BREAK.PAGE)

# ────────────────────────────────────────────────────────────────────────────

# ─── 6) Guardar el documento único en el directorio actual ──────────────────
master_doc.save("Todos_Los_Defectos.docx")
print("\n✅ Se generó 'Todos_Los_Defectos.docx' correctamente.")

